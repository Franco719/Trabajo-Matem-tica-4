#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import pandas as pd
import numpy as np
import re
from math import isnan

# ------------- Config -------------
INPUT_CSV = "csvbueno.csv"
# Parámetros de descenso del gradiente
ALPHA = 0.1
N_ITER = 433
# Nombre del docx de salida
DOCX_FILENAME = "regression_multivar_report.docx"
# -----------------------------------

# ----------------------------
# 1) Leer CSV probando encodings
# ----------------------------
df = None
last_exc = None
for enc in ('utf-8', 'latin-1', 'cp1252', 'ascii'):
    try:
        # algunos csv usan ; como separador; si no, pandas detectará
        df = pd.read_csv(INPUT_CSV, sep=';', engine='python', encoding=enc, skipinitialspace=True)
        print(f"Leído {INPUT_CSV} con encoding={enc!r} y sep=';'")
        break
    except Exception as e:
        last_exc = e
# si falló con ; probamos sin sep
if df is None:
    for enc in ('utf-8', 'latin-1', 'cp1252', 'ascii'):
        try:
            df = pd.read_csv(INPUT_CSV, encoding=enc, skipinitialspace=True)
            print(f"Leído {INPUT_CSV} con encoding={enc!r} y sep=',' (fallback)")
            break
        except Exception as e:
            last_exc = e

if df is None:
    raise last_exc

print(f"DataFrame shape: {df.shape}")
print("Columnas:", df.columns.tolist())

# ------------------------------------------------
# 2) Funciones de extracción / conversión
# ------------------------------------------------
def to_numeric_general(value):
    """Extrae números; si hay varios (rango), devuelve el promedio."""
    if pd.isna(value):
        return np.nan
    if isinstance(value, (int, float)):
        return float(value)
    s = str(value).strip()
    # buscamos todos los números (enteros o decimales) y promediamos
    nums = re.findall(r"[-+]?\d*[\.,]?\d+", s)
    if not nums:
        return np.nan
    try:
        nums_f = [float(x.replace(",", ".")) for x in nums]
        return float(sum(nums_f) / len(nums_f))
    except:
        return np.nan

def correct_hp_column(valor):
    """Extrae un número representativo de HorsePower (promedio si hay rango)."""
    return to_numeric_general(valor)

def extract_cc(valor):
    """
    Convierte CC textual a número en cc:
    - "1.2L" -> 1200
    - "1,200 cc" -> 1200
    - "3990 cc" -> 3990
    - Si detecta 'kw' o 'battery' devuelve NaN (eliminar eléctricos)
    - Si hay rango promedia
    """
    if pd.isna(valor):
        return np.nan
    s = str(valor).lower()
    if "kw" in s or "battery" in s:
        return np.nan
    # litros: "1.2 l" o "1.2l"
    m_l = re.search(r"(\d*[\.,]?\d+)\s*l\b", s)
    if m_l:
        try:
            litros = float(m_l.group(1).replace(",", "."))
            return litros * 1000.0
        except:
            pass
    # buscar ocurrencias con cc
    nums = re.findall(r"[-+]?\d*[\.,]?\d+", s)
    if not nums:
        return np.nan
    # convertimos y si detectamos que hay formatos de miles (con comas) quitamos comas
    nums_clean = []
    for x in nums:
        x_clean = x.replace(",", ".")
        try:
            nums_clean.append(float(x_clean))
        except:
            try:
                nums_clean.append(float(x.replace(",", "")))
            except:
                pass
    if not nums_clean:
        return np.nan
    # si hay más de 1 número (rango), promediamos; además si los números parecen pequeños (<50) y
    # no tenían L/cc, podríamos interpretarlos como litros -> multiplicar por 1000
    avg = sum(nums_clean) / len(nums_clean)
    # heurística: si avg < 50 -> puede estar en L (ej 1.2 -> 1.2 L), multiplicamos por 1000
    if avg < 50:
        # pero solo si en el texto aparece 'l' o no aparece 'cc'. Hacemos condición liviana:
        if 'l' in s and 'cc' not in s:
            return avg * 1000.0
        # si no hay 'cc' ni 'l' y avg < 50, preferimos devolver avg*1000? mejor no asumir: devolvemos avg*1000 solo si 'l' presente
    return avg

def extract_torque(valor):
    """Extrae número representativo de Torque (promedia rangos)."""
    return to_numeric_general(valor)

# ----------------------------
# 3) Aplicar limpieza a columnas
# ----------------------------
# Columnas esperadas
col_hp = "HorsePower"
col_cc = "CC/Battery Capacity"
col_torque = "Torque"
col_ts = "Total Speed"

for c in (col_hp, col_cc, col_torque, col_ts):
    if c not in df.columns:
        raise ValueError(f"La columna esperada '{c}' no se encuentra en el CSV. Columnas disponibles: {df.columns.tolist()}")

# Diagnóstico previo
#print("Valores no nulos antes de limpieza:", df[[col_hp, col_cc, col_torque, col_ts]].notna().sum().to_dict())

# Creamos mask_hp: filas inicialmente con HorsePower y Total Speed no nulos (para comparar y mantener Y)
mask_hp = df[col_hp].notna() & df[col_ts].notna()
n_hp = int(mask_hp.sum())
#print(f"Filas con HorsePower y Total Speed presentes (mask_hp): {n_hp}")

# Aplicar conversiones
df[col_hp] = df[col_hp].apply(correct_hp_column)
df[col_cc] = df[col_cc].apply(extract_cc)
df[col_torque] = df[col_torque].apply(extract_torque)
# Total Speed: extraer número (p. ej. "250 km/h" -> 250)
df[col_ts] = df[col_ts].astype(str).str.extract(r"([-+]?\d*[\.,]?\d+)")[0].replace("", np.nan).map(lambda v: float(str(v).replace(",", ".")) if pd.notna(v) else np.nan)

# Ver conteos después de extracción
#print("Valores no nulos después de extracción:", df[[col_hp, col_cc, col_torque, col_ts]].notna().sum().to_dict())

# ----------------------------
# 4) Imputación por mediana para CC y Torque entre las filas usadas por mask_hp (conservar n_hp)
# ----------------------------
cc_in_mask = df.loc[mask_hp, col_cc]
torque_in_mask = df.loc[mask_hp, col_torque]

cc_nonnull = int(cc_in_mask.notna().sum())
torque_nonnull = int(torque_in_mask.notna().sum())
cc_missing = n_hp - cc_nonnull
torque_missing = n_hp - torque_nonnull

#print(f"Entre las {n_hp} filas originales: CC numérico = {cc_nonnull}, faltantes = {cc_missing}")
#print(f"Entre las {n_hp} filas originales: Torque numérico = {torque_nonnull}, faltantes = {torque_missing}")

# imputar mediana calculada sobre las filas mask_hp que tengan valor
median_cc = cc_in_mask.median()
median_torque = torque_in_mask.median()
#print(f"Mediana CC (mask_hp): {median_cc}, Mediana Torque (mask_hp): {median_torque}")

# Rellenar en el DataFrame SOLO en las filas que están en mask_hp
df.loc[mask_hp, col_cc] = df.loc[mask_hp, col_cc].fillna(median_cc)
df.loc[mask_hp, col_torque] = df.loc[mask_hp, col_torque].fillna(median_torque)

# Ahora tomar solo las filas que pertenecen a mask_hp y que además tengan Total Speed no nulo y HorsePower no nulo
df_model = df.loc[mask_hp, [col_hp, col_cc, col_torque, col_ts]].dropna().reset_index(drop=True)
#print(f"Filas disponibles para el modelo (después de imputación y dropna final): {df_model.shape[0]}")

if df_model.shape[0] < 3:
    raise ValueError("Pocos datos para ajustar regresión múltiple (menos de 3 filas). Revisa limpieza.")

# ----------------------------
# 5) Preparar X, y y normalizar
# ----------------------------
X = df_model[[col_hp, col_cc, col_torque]].values.astype(float)  # orden x1,x2,x3
y = df_model[col_ts].values.astype(float)

# guardar medias y sds para desnormalizar
X_mean = X.mean(axis=0)
X_std = X.std(axis=0, ddof=0)
y_mean = y.mean()
y_std = y.std(ddof=0)

# evitar dividir por cero
X_std[X_std == 0] = 1.0
if y_std == 0:
    y_std = 1.0

X_norm = (X - X_mean) / X_std
y_norm = (y - y_mean) / y_std

# Agregar columna de 1s
X_b = np.c_[np.ones((X_norm.shape[0], 1)), X_norm]

# ----------------------------
# 6) Descenso del gradiente (vectorizado)
# ----------------------------
m = len(y_norm)
theta = np.zeros(X_b.shape[1], dtype=float)  # [b0_norm, b1_norm, b2_norm, b3_norm]

def compute_cost(Xb, yv, th):
    preds = Xb.dot(th)
    errs = preds - yv
    return (1.0 / (2*m)) * np.sum(errs**2)

def gradient_descent(Xb, yv, theta_init, alpha, n_iter):
    th = theta_init.copy()
    cost_hist = []
    for it in range(n_iter):
        grad = (1.0 / m) * Xb.T.dot(Xb.dot(th) - yv)
        th = th - alpha * grad
        cost_hist.append(compute_cost(Xb, yv, th))
    return th, cost_hist

theta_final, cost_history = gradient_descent(X_b, y_norm, theta, ALPHA, N_ITER)

# ----------------------------
# 7) Desnormalizar coeficientes a unidades reales
# ----------------------------
# theta_final corresponde a parámetros sobre variables normalizadas:
# y_norm = theta0 + theta1 * (x1-x1_mean)/x1_std + ...
# Para volver a unidades reales:
theta_real = np.zeros_like(theta_final)
theta_real[1:] = theta_final[1:] * (y_std / X_std)
theta_real[0] = y_mean - np.sum(theta_real[1:] * X_mean)

# nombres para informe
b0 = float(theta_real[0])
b1 = float(theta_real[1])
b2 = float(theta_real[2])
b3 = float(theta_real[3])

# predicciones y R2 en escala real
y_pred = b0 + b1 * X[:,0] + b2 * X[:,1] + b3 * X[:,2]
SST = np.sum((y - y_mean)**2)
SSR = np.sum((y_pred - y_mean)**2)
SSE = np.sum((y - y_pred)**2)
R2_model = SSR / SST if SST != 0 else np.nan

print("\n--- RESULTADOS REGRESIÓN MÚLTIPLE (Descenso del Gradiente) ---")
print(f"n = {m}")
print(f"Coeficientes (unidades reales):")
print(f"b0 = {b0:.6f}")
print(f"b1 (x1 HorsePower) = {b1:.6f}")
print(f"b2 (x2 CC) = {b2:.6f}")
print(f"b3 (x3 Torque) = {b3:.6f}")
print(f"Ecuación: Y = {b0:.4f} + {b1:.4f}*x1 + {b2:.4f}*x2 + {b3:.4f}*x3")
print(f"R² = {R2_model:.6f}")
print(f"SSE = {SSE:.4f}, SSR = {SSR:.4f}, SST = {SST:.4f}")
"""
# ----------------------------
# 8) Guardar informe Word (sin gráficos) - usando python-docx
# ----------------------------
try:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading('Reporte: Regresión Lineal Múltiple (Descenso del Gradiente)', level=1)
    doc.add_paragraph(f'Dataset: {INPUT_CSV}')
    doc.add_paragraph(f'Filas usadas (mask_hp originalmente): {n_hp}    Filas para modelo después limpieza: {df_model.shape[0]}')
    doc.add_paragraph('Variables:')
    doc.add_paragraph('Y = Velocidad Final (Total Speed)')
    doc.add_paragraph('x1 = HorsePower')
    doc.add_paragraph('x2 = CC/Battery Capacity (cc)')
    doc.add_paragraph('x3 = Torque')

    doc.add_heading('Coeficientes (unidades reales)', level=2)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Coeficiente'
    hdr_cells[1].text = 'Valor'
    for name, val in [('b0 (intercepto)', b0), ('b1 (x1 HorsePower)', b1), ('b2 (x2 CC)', b2), ('b3 (x3 Torque)', b3)]:
        r = table.add_row().cells
        r[0].text = name
        r[1].text = f"{val:.6f}"

    doc.add_paragraph('')
    doc.add_heading('Ecuación del modelo', level=2)
    doc.add_paragraph(f"Ŷ = {b0:.4f} + {b1:.4f}·x1 + {b2:.4f}·x2 + {b3:.4f}·x3")

    doc.add_heading('Medidas de ajuste', level=2)
    doc.add_paragraph(f"R² = {R2_model:.6f}")
    doc.add_paragraph(f"SST = {SST:.4f}")
    doc.add_paragraph(f"SSR = {SSR:.4f}")
    doc.add_paragraph(f"SSE = {SSE:.4f}")

    doc.add_heading('Análisis automático (breve)', level=2)
    # Generar análisis simple
    analisis = []
    analisis.append("El mejor predictor individual previamente observado fue HorsePower (x1).")
    if b1 > 0:
        analisis.append("El coeficiente b1 positivo indica que, manteniendo las demás variables constantes, un aumento de x1 tiende a aumentar la velocidad final.")
    else:
        analisis.append("El coeficiente b1 negativo/pequeño indica que el efecto de x1 no es claramente positivo en presencia de las otras variables (posible multicolinealidad).")
    analisis.append(f"El R² del modelo es {R2_model:.3f}, lo que indica la fracción de varianza de Y explicada por las tres variables.")
    # multicolinealidad simple: corrs
    corrs = df_model[[col_hp, col_cc, col_torque, col_ts]].corr()
    analisis.append("Se observan correlaciones importantes entre predictores, lo que sugiere multicolinealidad y posible inestabilidad de coeficientes.")
    for linea in analisis:
        doc.add_paragraph(linea)

    doc.add_paragraph('')
    doc.add_paragraph('NOTA: Se usó imputación por mediana para CC y Torque en las filas inicialmente con HorsePower y Total Speed presentes, para conservar n y permitir comparación con regresión simple.')

    doc.save(DOCX_FILENAME)
    print(f"\nInforme guardado en {DOCX_FILENAME}")
except Exception as e:
    print("\nNo se pudo generar el .docx automáticamente. Instala python-docx (pip install python-docx) si querés la salida en Word.")
    print("Error detalle:", e)
"""
# ----------------------------
# 9) (Opcional) imprimir la matriz de correlación para reporte
# ----------------------------
print("\nMatriz de correlación (variables usadas):")
print(df_model[[col_hp, col_cc, col_torque, col_ts]].corr().to_string())

# Fin del script